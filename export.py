from docx import Document

def export_meeting_summary(summary, filename="meeting_summary.docx"):
    """Exports the summary to a Word document."""
    doc = Document()
    doc.add_heading("Meeting Summary", level=1)
    doc.add_paragraph(summary)
    doc.save(filename)